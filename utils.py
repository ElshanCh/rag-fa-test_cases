from docx import Document
from langchain_core.documents.base import Document as dx
import json
import tiktoken


# --------------------------------------------------
# Function to extract tables
# --------------------------------------------------
def extract_format_tables(file_path):
    doc = Document(file_path)
    tables_list = []
    for table_index, table in enumerate(doc.tables):
        # print(f"Table {table_index + 1}:")
        table_info = {"table_name": f"Table {table_index + 1}", "table": ""}
        # table_rows = ""
        for row_index, row in enumerate(table.rows):
            row_text = []
            for cell_index, cell in enumerate(row.cells):
                cell_text = ' '.join(paragraph.text for paragraph in cell.paragraphs)
                row_text.append(f"Cell {cell_index + 1}: {cell_text}")

            # print(f"Row {row_index + 1}: {' | '.join(row_text)}")
            if row_index == 0:
                table_info["table"] = table_info["table"] + f"Row {row_index + 1}: {' | '.join(row_text)}"
            else:
                table_info["table"] = table_info["table"] + f"\nRow {row_index + 1}: {' | '.join(row_text)}"
        tables_list.append(table_info)
        
        # print(table_rows)
        # print("\n")
    return tables_list


# --------------------------------------------------
# Converting JSON in Document object
# --------------------------------------------------
def json_to_document(json_data):
    # Extract relevant fields from the JSON data
    # Replace the field names below with the actual field names in your JSON
    page_content = json_data.get("content", None)
    chunk_num = json_data.get("chunk_num", None)
    file_path = json_data.get("file_path", None)
    heading = json_data.get("heading", None)
    metadata={'source': file_path,'heading': heading, 'chunk_num': chunk_num}
    # ...

    # Create a Document object with the extracted fields
    document = dx(page_content=page_content, metadata=metadata)

    return document


# --------------------------------------------------
# Tiktoken to count the tokens
# --------------------------------------------------
def num_tokens_from_string(string: str, encoding_name: str) -> int:
    """Returns the number of tokens in a text string."""
    encoding = tiktoken.get_encoding(encoding_name)
    num_tokens = len(encoding.encode(string))
    return num_tokens


# --------------------------------------------------
# Doc Processor
# --------------------------------------------------
class DocxProcessor:
    """
    The `DocxProcessor` class encapsulates functionality for processing Microsoft Word documents (.docx files) by extracting 
    headings, paragraphs, and table information. This class provides methods to load headings, read content, and filter data 
    based on specified criteria. It is designed to enhance code modularity and readability when working with Word documents, 
    offering a convenient and organized approach to document processing.

    Methods:
    - __init__(self, file_path): Initializes a new instance of the `DocxProcessor` class with the specified file path.

    - load_headings(self): Extracts and stores headings from the Word document, identifying text formatted as headings.

    - read_content(self): Parses the Word document to retrieve information about paragraphs and tables, organizing the content 
    into a structured sequence.

    - filter_data(self): Filters the processed content, removing empty paragraphs and marking paragraphs containing specified 
    headings.

    - process_result(self): Processes the filtered data to generate a result based on specified conditions.

    Example Usage:
    ```python
    file_path = "your_document_path"
    doc_processor = DocxProcessor(file_path)
    doc_processor.load_headings()
    doc_processor.read_content()
    filtered_data = doc_processor.filter_data()
    result = doc_processor.process_result()
    print(result)
    ```
    """
    
    def __init__(self, file_path,tables_list):
        self.file_path = file_path
        self.headings = []
        self.content_sequence = []
        self.doc = Document(self.file_path)
        self.tables_list = tables_list

    def load_headings(self):
        
        for paragraph in self.doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                heading = paragraph.text
                self.headings.append({'heading': heading})

    def read_content(self):
        
        for element in self.doc.element.body:
            if element.tag.endswith('tbl'):
                # Handle tables
                table_info = {"type": "table", "content": ""}
                for row_index, row in enumerate(element):
                    row_text = []
                    for cell_index, cell in enumerate(row):
                        cell_text = ""
                        for paragraph in cell:
                            if paragraph.text is not None:
                                cell_text += paragraph.text
                        row_text.append(f"Cell {cell_index + 1}: {cell_text}")

                    if row_index == 0:
                        table_info["content"] += f"Row {row_index + 1}: {' | '.join(row_text)}"
                    else:
                        table_info["content"] += f"\nRow {row_index + 1}: {' | '.join(row_text)}"

                self.content_sequence.append(table_info)
            elif element.tag.endswith('p'):
                # Handle paragraphs
                paragraph_text = element.text
                self.content_sequence.append({'type': 'paragraph', 'content': paragraph_text})
        return self.content_sequence

    def filter_data(self):
        filtered_data = [item for item in self.content_sequence if not (item.get('type') == 'paragraph' and not item.get('content'))]
        headings_list = [heading["heading"] for heading in self.headings]

        for item in filtered_data:
            if item['type'] == "paragraph" and item['content'] in headings_list:
                item['heading'] = True
            else:
                item['heading'] = False

        return filtered_data
    
    def chunking(self):
        result = []
        current_heading = None
        current_chunk = ""
        chunk_count = 1
        table_num = 0
        extracted_tables = [item for item in self.filter_data() if item.get("type") == "table"]
        for item in self.filter_data():
            if item["type"] == "table" and len(extracted_tables) == len(extracted_tables):
                item["content"] = self.tables_list[table_num]["table"]
                table_num +=1
            if item['heading']:
                if current_heading is None and len(current_chunk) > 0:
                    result.append({"heading": None, "content": f"Header: None\n{current_chunk}","chunk_num": chunk_count, "file_path": self.file_path})
                    chunk_count += 1

                if current_heading is not None:
                    result.append({"heading": current_heading, "content": f"Header: {current_heading}\n{current_chunk}","chunk_num": chunk_count, "file_path": self.file_path})
                    chunk_count += 1

                current_chunk = ""
                current_heading = item['content']
            else:
                current_chunk += item['content'] + "\n"

        # Add the last chunk to the result
        if current_heading is not None:
            result.append({"heading": current_heading, "content": f"Header: {current_heading}\n{current_chunk}","chunk_num": chunk_count, "file_path": self.file_path})
        else:
            # If there are no headings with False, create an entry with an empty content
            result.append({"heading": None, "content": f"Header: None\n{current_chunk}","chunk_num": chunk_count, "file_path": self.file_path})

        return result
    




